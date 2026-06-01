"""Word Agent — creates professionally formatted Word (.docx) documents.

Generates cover page, auto TOC, styled sections, headers/footers, and tables.
Uses python-docx with polished styles and consistent visual identity.
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from future_agents.core.base_agent import BaseAgent, TaskContext, TaskResult
from future_agents.models.feedback import ExecutionOutcome

logger = logging.getLogger(__name__)

# Professional color palette for Word docs (hex strings as used by docx)
WORD_THEMES = {
    "corporate_blue": {
        "heading1": "1565C0",
        "heading2": "1976D2",
        "heading3": "42A5F5",
        "accent": "1565C0",
        "cover_bg": "1565C0",
        "cover_text": "FFFFFF",
        "header_bg": "E3F2FD",
        "table_header": "1565C0",
    },
    "executive_green": {
        "heading1": "2E7D32",
        "heading2": "388E3C",
        "heading3": "66BB6A",
        "accent": "2E7D32",
        "cover_bg": "1B5E20",
        "cover_text": "FFFFFF",
        "header_bg": "E8F5E9",
        "table_header": "2E7D32",
    },
    "modern_purple": {
        "heading1": "4A148C",
        "heading2": "7B1FA2",
        "heading3": "AB47BC",
        "accent": "7B1FA2",
        "cover_bg": "4A148C",
        "cover_text": "FFFFFF",
        "header_bg": "F3E5F5",
        "table_header": "7B1FA2",
    },
    "classic_navy": {
        "heading1": "0D47A1",
        "heading2": "1565C0",
        "heading3": "5C6BC0",
        "accent": "0D47A1",
        "cover_bg": "0D3B6E",
        "cover_text": "F5F5F5",
        "header_bg": "E8EAF6",
        "table_header": "0D47A1",
    },
}


class WordAgent(BaseAgent):
    """Creates professional Word documents with polished formatting.

    Capabilities:
    - word.create  — full document from title + sections
    - word.export  — confirm file exists and return metadata
    """

    def __init__(self, output_dir: str = "output/word", **kwargs: Any) -> None:
        super().__init__(**kwargs)
        self._output_dir = Path(output_dir)
        self._output_dir.mkdir(parents=True, exist_ok=True)

    @property
    def agent_type(self) -> str:
        return "word"

    @property
    def capabilities(self) -> list[str]:
        return ["word.create", "word.export"]

    async def _execute(self, context: TaskContext) -> TaskResult:
        handlers = {
            "word.create": self._create_document,
            "word.export": self._export_document,
        }
        handler = handlers.get(context.intent)
        if not handler:
            return TaskResult(
                task_id=context.task_id,
                agent_id=self.agent_id,
                outcome=ExecutionOutcome.FAILURE,
                errors=[f"Unknown intent: {context.intent}"],
            )
        return await handler(context)

    # ── Helpers ──────────────────────────────────────────────────────────────

    @staticmethod
    def _set_cell_color(cell, hex_color: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    @staticmethod
    def _add_horizontal_rule(doc) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        para = doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _style_heading(self, doc, level: int, theme: dict) -> None:
        from docx.shared import Pt, RGBColor

        style_name = f"Heading {level}"
        style = doc.styles[style_name]
        font = style.font
        color_hex = theme.get(f"heading{level}", "000000")
        r = int(color_hex[0:2], 16)
        g = int(color_hex[2:4], 16)
        b = int(color_hex[4:6], 16)
        font.color.rgb = RGBColor(r, g, b)
        font.bold = True
        if level == 1:
            font.size = Pt(24)
        elif level == 2:
            font.size = Pt(18)
        else:
            font.size = Pt(14)

    # ── Core document creation ────────────────────────────────────────────────

    async def _create_document(self, context: TaskContext) -> TaskResult:
        """Build a complete professional Word document.

        Parameters:
            title         (str): Document title
            subtitle      (str): Optional subtitle
            author        (str): Author name
            organization  (str): Organization / company name
            color_theme   (str): corporate_blue|executive_green|modern_purple|classic_navy
            sections      (list[dict]): Each has:
                            - heading (str)
                            - content (str)
                            - subsections (list[dict]): Optional {heading, content}
                            - table (dict): Optional {headers: [], rows: [[]]}
                            - bullets (list[str]): Optional bullet list
            summary       (str): Closing summary paragraph
            output_name   (str): Filename without extension
        """
        try:
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Inches, Pt, RGBColor
        except ImportError:
            return TaskResult(
                task_id=context.task_id,
                agent_id=self.agent_id,
                outcome=ExecutionOutcome.FAILURE,
                errors=["python-docx not installed. Run: pip install python-docx"],
            )

        p = context.parameters
        title = p.get("title", "Professional Document")
        subtitle = p.get("subtitle", "")
        author = p.get("author", "Document Agent")
        organization = p.get("organization", "")
        theme_name = p.get("color_theme", "corporate_blue")
        sections = p.get("sections", [])
        summary_text = p.get("summary", "")
        output_name = p.get("output_name", title.lower().replace(" ", "_")[:40])

        theme = WORD_THEMES.get(theme_name, WORD_THEMES["corporate_blue"])

        doc = Document()

        # ── Page margins ────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # ── Apply heading styles ─────────────────────────────────────
        for level in (1, 2, 3):
            try:
                self._style_heading(doc, level, theme)
            except Exception:
                pass

        # ── Cover page ──────────────────────────────────────────────
        accent_hex = theme["cover_bg"]
        ar = int(accent_hex[0:2], 16)
        ag = int(accent_hex[2:4], 16)
        ab = int(accent_hex[4:6], 16)

        cover_heading = doc.add_heading("", level=0)
        cover_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_run = cover_heading.add_run(title)
        cover_run.font.size = Pt(36)
        cover_run.font.bold = True
        cover_run.font.color.rgb = RGBColor(ar, ag, ab)

        if subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.size = Pt(18)
            sub_run.font.italic = True
            sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_paragraph()

        # Cover metadata table
        meta_table = doc.add_table(rows=3, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_data = [
            ("Author", author),
            ("Organization", organization or "—"),
            ("Date", datetime.now().strftime("%B %d, %Y")),
        ]
        for i, (label, value) in enumerate(meta_data):
            row = meta_table.rows[i]
            lbl_cell = row.cells[0]
            val_cell = row.cells[1]
            lbl_cell.text = label
            val_cell.text = value
            lbl_cell.paragraphs[0].runs[0].font.bold = True
            lbl_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(ar, ag, ab)

        doc.add_page_break()

        # ── Header / Footer ──────────────────────────────────────────
        for sec in doc.sections:
            header = sec.header
            footer = sec.footer

            hdr_para = header.paragraphs[0]
            hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_run = hdr_para.add_run(f"  {title}  ")
            hdr_run.font.size = Pt(9)
            hdr_run.font.italic = True
            hdr_run.font.color.rgb = RGBColor(ar, ag, ab)

            ftr_para = footer.paragraphs[0]
            ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ftr_para.add_run(f"{organization}  |  {author}  |  Page ").font.size = Pt(8)
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), "begin")
            ftr_para._p.append(fld)
            instrText = OxmlElement("w:instrText")
            instrText.text = "PAGE"
            ftr_para._p.append(instrText)
            fld2 = OxmlElement("w:fldChar")
            fld2.set(qn("w:fldCharType"), "end")
            ftr_para._p.append(fld2)

        # ── TOC placeholder ──────────────────────────────────────────
        toc_heading = doc.add_heading("Table of Contents", level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Word TOC field
        toc_para = doc.add_paragraph()
        from docx.oxml import OxmlElement as OXE
        from docx.oxml.ns import qn as QN

        fld_begin = OXE("w:fldChar")
        fld_begin.set(QN("w:fldCharType"), "begin")
        fld_begin.set(QN("w:dirty"), "true")
        toc_para._p.append(fld_begin)
        instr = OXE("w:instrText")
        instr.set(QN("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        toc_para._p.append(instr)
        fld_end = OXE("w:fldChar")
        fld_end.set(QN("w:fldCharType"), "end")
        toc_para._p.append(fld_end)

        toc_note = doc.add_paragraph()
        run = toc_note.add_run("(Update table of contents: press Ctrl+A then F9 in Word)")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_page_break()

        # ── Sections ─────────────────────────────────────────────────
        section_count = 0
        for sec_data in sections:
            section_count += 1
            heading = sec_data.get("heading", f"Section {section_count}")
            content = sec_data.get("content", "")

            doc.add_heading(f"{section_count}. {heading}", level=1)

            if content:
                content_para = doc.add_paragraph(content)
                content_para.paragraph_format.space_after = Pt(6)
                content_para.paragraph_format.line_spacing = Pt(15)

            # Subsections
            subsections = sec_data.get("subsections", [])
            for j, sub in enumerate(subsections):
                sub_heading = sub.get("heading", f"Subsection {j + 1}")
                sub_content = sub.get("content", "")
                doc.add_heading(f"{section_count}.{j + 1} {sub_heading}", level=2)
                if sub_content:
                    sp = doc.add_paragraph(sub_content)
                    sp.paragraph_format.line_spacing = Pt(14)

            # Bullet list
            bullets = sec_data.get("bullets", [])
            for bullet in bullets:
                bp = doc.add_paragraph(style="List Bullet")
                run = bp.add_run(bullet)
                run.font.size = Pt(11)

            # Table
            table_data = sec_data.get("table")
            if table_data:
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                if headers:
                    cols = len(headers)
                    tbl = doc.add_table(rows=1 + len(rows), cols=cols)
                    tbl.style = "Table Grid"

                    # Header row
                    hdr_row = tbl.rows[0]
                    th_hex = theme["table_header"]
                    int(th_hex[0:2], 16)
                    int(th_hex[2:4], 16)
                    int(th_hex[4:6], 16)
                    for ci, hdr_text in enumerate(headers):
                        cell = hdr_row.cells[ci]
                        cell.text = hdr_text
                        self._set_cell_color(cell, th_hex)
                        run = cell.paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(10)

                    # Data rows
                    for ri, row_data in enumerate(rows):
                        data_row = tbl.rows[ri + 1]
                        for ci, cell_val in enumerate(row_data):
                            cell = data_row.cells[ci]
                            cell.text = str(cell_val)
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                            if ri % 2 == 0:
                                self._set_cell_color(cell, "F5F5F5")

                    doc.add_paragraph()

            self._add_horizontal_rule(doc)

        # ── Summary ──────────────────────────────────────────────────
        if summary_text:
            doc.add_heading("Summary", level=1)
            s_para = doc.add_paragraph(summary_text)
            s_para.paragraph_format.line_spacing = Pt(15)

        # ── Save ─────────────────────────────────────────────────────
        out_path = self._output_dir / f"{output_name}.docx"
        doc.save(str(out_path))

        size_kb = out_path.stat().st_size // 1024

        logger.info("Word doc saved: %s (%d KB, %d sections)", out_path, size_kb, section_count)
        await self.emit("word.created", {"path": str(out_path), "sections": section_count})

        return TaskResult(
            task_id=context.task_id,
            agent_id=self.agent_id,
            outcome=ExecutionOutcome.SUCCESS,
            data={
                "file_path": str(out_path),
                "size_kb": size_kb,
                "section_count": section_count,
                "title": title,
                "theme": theme_name,
                "structure": ["Cover Page", "Table of Contents"]
                + [s.get("heading", "Section") for s in sections]
                + (["Summary"] if summary_text else []),
            },
        )

    async def _export_document(self, context: TaskContext) -> TaskResult:
        file_path = context.parameters.get("file_path", "")
        path = Path(file_path)
        if not path.exists():
            return TaskResult(
                task_id=context.task_id,
                agent_id=self.agent_id,
                outcome=ExecutionOutcome.FAILURE,
                errors=[f"File not found: {file_path}"],
            )
        return TaskResult(
            task_id=context.task_id,
            agent_id=self.agent_id,
            outcome=ExecutionOutcome.SUCCESS,
            data={"file_path": str(path), "size_bytes": path.stat().st_size, "ready": True},
        )

    async def assess_self(self) -> dict[str, Any]:
        files = list(self._output_dir.glob("*.docx"))
        return {
            "output_dir": str(self._output_dir),
            "documents_created": len(files),
            "available_themes": list(WORD_THEMES.keys()),
        }
